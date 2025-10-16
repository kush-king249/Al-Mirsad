#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Report Generation Module for Al-Mirsad
وحدة إنشاء التقارير لأداة المرصاد

This module handles generating comprehensive incident response reports.
تتعامل هذه الوحدة مع إنشاء تقارير شاملة للاستجابة للحوادث.
"""

import os
import json
import datetime
import logging
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


class ReportGenerator:
    """
    Report generation class for creating incident response reports.
    فئة إنشاء التقارير لإنشاء تقارير الاستجابة للحوادث.
    """
    
    def __init__(self, output_dir: str = "reports"):
        """
        Initialize the ReportGenerator.
        
        Args:
            output_dir (str): Directory to store generated reports
        """
        self.output_dir = output_dir
        self.logger = self._setup_logger()
        self._ensure_output_dir()
    
    def _setup_logger(self) -> logging.Logger:
        """
        Setup logging for the module.
        """
        logger = logging.getLogger(__name__)
        logger.setLevel(logging.INFO)
        
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter(
                '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
            )
            handler.setFormatter(formatter)
            logger.addHandler(handler)
        
        return logger
    
    def _ensure_output_dir(self):
        """
        Ensure output directory exists.
        """
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            self.logger.info(f"Created output directory: {self.output_dir}")
    
    def generate_incident_report(self, incident_data: Dict[str, Any], 
                               format_type: str = "docx") -> str:
        """
        Generate a comprehensive incident response report.
        إنشاء تقرير شامل للاستجابة للحوادث.
        
        Args:
            incident_data (Dict[str, Any]): Incident data and analysis results
            format_type (str): Output format ('docx', 'pdf', or 'json')
            
        Returns:
            str: Path to the generated report file
        """
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        incident_id = incident_data.get('incident_id', f'INC_{timestamp}')
        
        if format_type.lower() == 'docx':
            return self._generate_docx_report(incident_data, incident_id, timestamp)
        elif format_type.lower() == 'pdf':
            return self._generate_pdf_report(incident_data, incident_id, timestamp)
        elif format_type.lower() == 'json':
            return self._generate_json_report(incident_data, incident_id, timestamp)
        else:
            self.logger.error(f"Unsupported format type: {format_type}")
            return ""
    
    def _generate_docx_report(self, incident_data: Dict[str, Any], 
                            incident_id: str, timestamp: str) -> str:
        """Generate DOCX format report."""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('تقرير الاستجابة للحوادث الأمنية', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Incident Response Report', 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Incident Overview
            doc.add_heading('معلومات الحادث - Incident Information', 1)
            
            # Create table for incident details
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Field / الحقل'
            hdr_cells[1].text = 'Value / القيمة'
            
            # Add incident details
            incident_details = [
                ('Incident ID / معرف الحادث', incident_id),
                ('Report Date / تاريخ التقرير', datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Incident Type / نوع الحادث', incident_data.get('incident_type', 'Unknown')),
                ('Severity / الخطورة', incident_data.get('severity', 'Unknown')),
                ('Status / الحالة', incident_data.get('status', 'Under Investigation')),
                ('Affected Systems / الأنظمة المتأثرة', str(len(incident_data.get('affected_systems', [])))),
            ]
            
            for field, value in incident_details:
                row_cells = table.add_row().cells
                row_cells[0].text = field
                row_cells[1].text = str(value)
            
            # Executive Summary
            doc.add_heading('الملخص التنفيذي - Executive Summary', 1)
            summary_text = incident_data.get('executive_summary',
                'This report documents the incident response activities and findings.')
            doc.add_paragraph(summary_text)
            
            # Timeline
            doc.add_heading('الجدول الزمني - Timeline', 1)
            timeline = incident_data.get('timeline', [])
            if timeline:
                for event in timeline:
                    p = doc.add_paragraph()
                    p.add_run(f"{event.get('timestamp', '')}: ").bold = True
                    p.add_run(event.get('description', ''))
            else:
                doc.add_paragraph('No timeline data available.')
            
            # Log Analysis Results
            if 'log_analysis' in incident_data:
                doc.add_heading('تحليل السجلات - Log Analysis', 1)
                log_data = incident_data['log_analysis']
                
                for log_type, entries in log_data.items():
                    doc.add_heading(f'{log_type.title()} Logs', 2)
                    if entries:
                        for entry in entries[:10]:  # Limit to first 10 entries
                            doc.add_paragraph(entry, style='List Bullet')
                    else:
                        doc.add_paragraph('No suspicious entries found.')
            
            # Malware Analysis Results
            if 'malware_analysis' in incident_data:
                doc.add_heading('تحليل البرمجيات الخبيثة - Malware Analysis', 1)
                malware_data = incident_data['malware_analysis']
                
                if isinstance(malware_data, list):
                    for analysis in malware_data:
                        doc.add_heading(f"File: {analysis.get('file_name', 'Unknown')}", 2)
                        
                        # Risk score
                        risk_score = analysis.get('risk_score', 0)
                        p = doc.add_paragraph()
                        p.add_run('Risk Score / درجة المخاطر: ').bold = True
                        p.add_run(f'{risk_score}/100')
                        
                        # Threat classification
                        threat = analysis.get('threat_classification', {})
                        if threat.get('threat_type') != 'Unknown':
                            p = doc.add_paragraph()
                            p.add_run('Threat Type / نوع التهديد: ').bold = True
                            p.add_run(threat.get('threat_type', 'Unknown'))
                        
                        # Indicators
                        indicators = threat.get('indicators', [])
                        if indicators:
                            doc.add_paragraph('Indicators / المؤشرات:', style='Heading 3')
                            for indicator in indicators[:5]:  # Limit to first 5
                                doc.add_paragraph(indicator, style='List Bullet')
            
            # Network Isolation Status
            if 'network_isolation' in incident_data:
                doc.add_heading('حالة عزل الشبكة - Network Isolation Status', 1)
                isolation_data = incident_data['network_isolation']
                
                for system, status in isolation_data.items():
                    p = doc.add_paragraph()
                    p.add_run(f'{system}: ').bold = True
                    p.add_run(status)
            
            # Recommendations
            doc.add_heading('التوصيات - Recommendations', 1)
            recommendations = incident_data.get('recommendations', [
                'Continue monitoring affected systems',
                'Update security policies and procedures',
                'Conduct security awareness training'
            ])
            
            for i, recommendation in enumerate(recommendations, 1):
                doc.add_paragraph(f'{i}. {recommendation}')
            
            # Conclusion
            doc.add_heading('الخلاصة - Conclusion', 1)
            conclusion = incident_data.get('conclusion',
                'The incident has been contained and is under investigation. '
                'All affected systems have been isolated and analyzed.')
            doc.add_paragraph(conclusion)
            
            # Footer
            doc.add_paragraph('\n' + '='*50)
            footer_p = doc.add_paragraph()
            footer_p.add_run('Generated by Al-Mirsad Incident Response Tool\n').italic = True
            footer_p.add_run(f'Report generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
            
            # Save document
            filename = f'incident_report_{incident_id}_{timestamp}.docx'
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            self.logger.info(f'DOCX report generated: {filepath}')
            return filepath
            
        except Exception as e:
            self.logger.error(f'Error generating DOCX report: {str(e)}')
            return ""
    
    def _generate_pdf_report(self, incident_data: Dict[str, Any], 
                           incident_id: str, timestamp: str) -> str:
        """Generate PDF format report."""
        try:
            class ArabicPDF(FPDF):
                def header(self):
                    self.set_font('Arial', 'B', 15)
                    self.cell(0, 10, 'Al-Mirsad Incident Response Report', 0, 1, 'C')
                    self.ln(10)
                
                def footer(self):
                    self.set_y(-15)
                    self.set_font('Arial', 'I', 8)
                    self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
            
            pdf = ArabicPDF()
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            
            # Title
            pdf.cell(0, 10, 'Incident Response Report', 0, 1, 'C')
            pdf.ln(10)
            
            # Incident Information
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, 'Incident Information', 0, 1)
            pdf.set_font('Arial', '', 12)
            
            incident_info = [
                f'Incident ID: {incident_id}',
                f'Report Date: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
                f'Incident Type: {incident_data.get("incident_type", "Unknown")}',
                f'Severity: {incident_data.get("severity", "Unknown")}',
                f'Status: {incident_data.get("status", "Under Investigation")}',
                f'Affected Systems: {len(incident_data.get("affected_systems", []))}'
            ]
            
            for info in incident_info:
                pdf.cell(0, 8, info, 0, 1)
            
            pdf.ln(5)
            
            # Executive Summary
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, 'Executive Summary', 0, 1)
            pdf.set_font('Arial', '', 12)
            
            summary = incident_data.get('executive_summary',
                'This report documents the incident response activities and findings.')
            
            # Split long text into multiple lines
            summary_lines = [summary[i:i+80] for i in range(0, len(summary), 80)]
            for line in summary_lines:
                pdf.cell(0, 6, line, 0, 1)
            
            pdf.ln(5)
            
            # Log Analysis Summary
            if 'log_analysis' in incident_data:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, 'Log Analysis Summary', 0, 1)
                pdf.set_font('Arial', '', 12)
                
                log_data = incident_data['log_analysis']
                for log_type, entries in log_data.items():
                    pdf.cell(0, 8, f'{log_type.title()} Logs: {len(entries)} suspicious entries', 0, 1)
                
                pdf.ln(5)
            
            # Malware Analysis Summary
            if 'malware_analysis' in incident_data:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, 'Malware Analysis Summary', 0, 1)
                pdf.set_font('Arial', '', 12)
                
                malware_data = incident_data['malware_analysis']
                if isinstance(malware_data, list):
                    for analysis in malware_data:
                        file_name = analysis.get('file_name', 'Unknown')
                        risk_score = analysis.get('risk_score', 0)
                        threat_type = analysis.get('threat_classification', {}).get('threat_type', 'Unknown')
                        
                        pdf.cell(0, 6, f'File: {file_name}', 0, 1)
                        pdf.cell(0, 6, f'  Risk Score: {risk_score}/100', 0, 1)
                        pdf.cell(0, 6, f'  Threat Type: {threat_type}', 0, 1)
                        pdf.ln(2)
            
            # Recommendations
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, 'Recommendations', 0, 1)
            pdf.set_font('Arial', '', 12)
            
            recommendations = incident_data.get('recommendations', [
                'Continue monitoring affected systems',
                'Update security policies and procedures',
                'Conduct security awareness training'
            ])
            
            for i, recommendation in enumerate(recommendations, 1):
                pdf.cell(0, 6, f'{i}. {recommendation}', 0, 1)
            
            # Save PDF
            filename = f'incident_report_{incident_id}_{timestamp}.pdf'
            filepath = os.path.join(self.output_dir, filename)
            pdf.output(filepath)
            
            self.logger.info(f'PDF report generated: {filepath}')
            return filepath
            
        except Exception as e:
            self.logger.error(f'Error generating PDF report: {str(e)}')
            return ""
    
    def _generate_json_report(self, incident_data: Dict[str, Any], 
                            incident_id: str, timestamp: str) -> str:
        """Generate JSON format report."""
        try:
            # Create comprehensive JSON report
            json_report = {
                'report_metadata': {
                    'incident_id': incident_id,
                    'report_generated': datetime.datetime.now().isoformat(),
                    'generator': 'Al-Mirsad Incident Response Tool',
                    'version': '1.0.0'
                },
                'incident_data': incident_data
            }
            
            # Save JSON
            filename = f'incident_report_{incident_id}_{timestamp}.json'
            filepath = os.path.join(self.output_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(json_report, f, indent=2, ensure_ascii=False, default=str)
            
            self.logger.info(f'JSON report generated: {filepath}')
            return filepath
            
        except Exception as e:
            self.logger.error(f'Error generating JSON report: {str(e)}')
            return ""
    
    def generate_summary_dashboard(self, incidents: List[Dict[str, Any]]) -> str:
        """
        Generate a summary dashboard report for multiple incidents.
        إنشاء تقرير لوحة معلومات موجزة لحوادث متعددة.
        
        Args:
            incidents (List[Dict[str, Any]]): List of incident data
            
        Returns:
            str: Path to the generated dashboard report
        """
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('لوحة معلومات الحوادث الأمنية', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading('Security Incidents Dashboard', 1)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Summary Statistics
            doc.add_heading('إحصائيات عامة - Summary Statistics', 1)
            
            total_incidents = len(incidents)
            high_severity = sum(1 for inc in incidents if inc.get('severity') == 'High')
            medium_severity = sum(1 for inc in incidents if inc.get('severity') == 'Medium')
            low_severity = sum(1 for inc in incidents if inc.get('severity') == 'Low')
            
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Table Grid'
            
            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = 'Metric / المقياس'
            hdr_cells[1].text = 'Value / القيمة'
            
            stats_data = [
                ('Total Incidents / إجمالي الحوادث', str(total_incidents)),
                ('High Severity / خطورة عالية', str(high_severity)),
                ('Medium Severity / خطورة متوسطة', str(medium_severity)),
                ('Low Severity / خطورة منخفضة', str(low_severity))
            ]
            
            for metric, value in stats_data:
                row_cells = stats_table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = value
            
            # Recent Incidents
            doc.add_heading('الحوادث الأخيرة - Recent Incidents', 1)
            
            if incidents:
                incidents_table = doc.add_table(rows=1, cols=4)
                incidents_table.style = 'Table Grid'
                
                hdr_cells = incidents_table.rows[0].cells
                hdr_cells[0].text = 'Incident ID'
                hdr_cells[1].text = 'Type / النوع'
                hdr_cells[2].text = 'Severity / الخطورة'
                hdr_cells[3].text = 'Status / الحالة'
                
                for incident in incidents[:10]:  # Show last 10 incidents
                    row_cells = incidents_table.add_row().cells
                    row_cells[0].text = incident.get('incident_id', 'Unknown')
                    row_cells[1].text = incident.get('incident_type', 'Unknown')
                    row_cells[2].text = incident.get('severity', 'Unknown')
                    row_cells[3].text = incident.get('status', 'Unknown')
            
            # Generate filename and save
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f'incidents_dashboard_{timestamp}.docx'
            filepath = os.path.join(self.output_dir, filename)
            doc.save(filepath)
            
            self.logger.info(f'Dashboard report generated: {filepath}')
            return filepath
            
        except Exception as e:
            self.logger.error(f'Error generating dashboard report: {str(e)}')
            return ""
    
    def create_incident_template(self) -> Dict[str, Any]:
        """
        Create a template for incident data structure.
        إنشاء قالب لهيكل بيانات الحادث.
        
        Returns:
            Dict[str, Any]: Incident data template
        """
        template = {
            'incident_id': '',
            'incident_type': '',  # e.g., 'Malware', 'Data Breach', 'DDoS'
            'severity': '',  # 'High', 'Medium', 'Low'
            'status': 'Under Investigation',  # 'Open', 'Under Investigation', 'Contained', 'Closed'
            'discovery_date': '',
            'affected_systems': [],
            'executive_summary': '',
            'timeline': [
                {
                    'timestamp': '',
                    'description': '',
                    'action_taken': ''
                }
            ],
            'log_analysis': {
                'system': [],
                'security': [],
                'application': []
            },
            'malware_analysis': [],
            'network_isolation': {},
            'recommendations': [],
            'conclusion': '',
            'analyst': '',
            'contact_info': ''
        }
        
        return template

